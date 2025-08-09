# backend/resume_parser.py

from io import BytesIO
from typing import Union
from PyPDF2 import PdfReader
from docx import Document

def parse_resume(file: Union[BytesIO, str], filename: str = None) -> str:
    """
    Parse uploaded resume file (.pdf or .docx) and return extracted plain text.
    
    Args:
        file: Uploaded file as BytesIO or path string.
        filename: Filename string to infer file type if needed.
        
    Returns:
        Extracted text as a single string.
    """
    if filename is None and isinstance(file, str):
        filename = file

    if filename is None:
        raise ValueError("Filename must be provided if file is not BytesIO.")

    ext = filename.split('.')[-1].lower()

    if ext == 'pdf':
        return _parse_pdf(file)
    elif ext == 'docx':
        return _parse_docx(file)
    else:
        raise ValueError(f"Unsupported file format: {ext}")

def _parse_pdf(file: Union[BytesIO, str]) -> str:
    text = []
    reader = PdfReader(file)
    for page in reader.pages:
        text.append(page.extract_text() or "")
    return "\n".join(text)

def _parse_docx(file: Union[BytesIO, str]) -> str:
    if isinstance(file, BytesIO):
        document = Document(file)
    else:
        document = Document(file)
    paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
    return "\n".join(paragraphs)
